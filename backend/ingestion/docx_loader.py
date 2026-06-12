import logging
from docx import Document
from .base_loader import DocumentLoader

logger = logging.getLogger(__name__)

class DocxLoader(DocumentLoader):
    """Loader for Word (DOCX) documents."""

    def extract_text(self, file_path: str) -> str:
        """
        Extract text from a DOCX file using python-docx.
        
        Args:
            file_path (str): Path to the DOCX file.
            
        Returns:
            str: The extracted text.
        """
        logger.info(f"Extracting text from DOCX: {file_path}")
        try:
            doc = Document(file_path)
            text_chunks = []
            
            # Extract from paragraphs
            for para in doc.paragraphs:
                if para.text:
                    text_chunks.append(para.text)
                    
            # Extract from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text:
                            text_chunks.append(cell.text)
                            
            return "\n".join(text_chunks)
        except Exception as e:
            logger.error(f"Error extracting text from DOCX {file_path}: {e}")
            raise
